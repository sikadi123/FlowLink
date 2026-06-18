from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(r"C:\Users\sikad\Desktop\FlowLink")
OUT = ROOT / "outputs" / "FlowLink即时通讯系统项目验收答辩发言稿_新版.docx"

GREEN = RGBColor(7, 193, 96)
DARK_GREEN = RGBColor(5, 120, 75)
INK = RGBColor(22, 32, 42)
MUTED = RGBColor(102, 112, 133)
BLUE = RGBColor(47, 128, 237)
SOFT = "F4F7FA"
MINT = "E8F7EF"
BLUE_SOFT = "E8F1FF"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", size=11, bold=False, color=INK, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(16 if level == 1 else 13)
    r.font.color.rgb = DARK_GREEN if level == 1 else BLUE
    r.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        r = p.add_run("• ")
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.color.rgb = GREEN
        r.bold = True
        r2 = p.add_run(item)
        r2.font.name = "Microsoft YaHei"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = INK


def add_callout(doc, title, body, fill=MINT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.margin_top = Inches(0.08)
    cell.margin_bottom = Inches(0.08)
    cell.margin_left = Inches(0.12)
    cell.margin_right = Inches(0.12)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    rt = p.add_run(title)
    rt.font.name = "Microsoft YaHei"
    rt._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rt.bold = True
    rt.font.size = Pt(11.5)
    rt.font.color.rgb = DARK_GREEN
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    rb = p2.add_run(body)
    rb.font.name = "Microsoft YaHei"
    rb._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rb.font.size = Pt(10.5)
    rb.font.color.rgb = INK
    add_para(doc, "", after=2)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("FlowLink 即时通讯系统项目验收答辩发言稿")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(22)
r.font.color.rgb = INK
r.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(16)
rs = sub.add_run("项目小组汇报稿 · 可配合新版 PPT 使用 · 2026 年 6 月")
rs.font.name = "Microsoft YaHei"
rs._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
rs.font.size = Pt(10.5)
rs.font.color.rgb = MUTED

add_callout(
    doc,
    "汇报定位",
    "本稿按“项目背景—架构设计—核心功能—新增亮点—测试验证—总结反思”的顺序组织，可用于 8 到 12 分钟课程验收答辩。"
)

add_heading(doc, "一、开场介绍")
add_para(doc, "各位老师、同学大家好，我们小组本次汇报的项目是 FlowLink 即时通讯系统。这个项目最初的目标是做一个在线聊天平台 Demo，但在后续迭代中，我们不再只满足于简单的消息收发，而是把它逐步扩展成一个具备账号体系、私聊群聊、群管理、通知中心、文件与语音能力、局域网视频通话、AI 助手、移动端 APK 和个性化设置的综合即时通讯平台。")
add_para(doc, "今天的汇报会重点说明三件事：第一，我们如何按照需求分析和设计方案完成系统架构；第二，目前项目已经具备哪些可演示功能；第三，在开发过程中遇到了哪些技术难点，以及我们对项目管理和后续优化的反思。")

add_heading(doc, "二、项目背景与目标")
add_para(doc, "FlowLink 的核心定位是“面向课程验收的在线即时通讯平台”。它既要展示即时通讯系统的基本业务能力，也要体现工程实现的完整性。因此，我们在需求层面把目标拆成四类：")
add_bullets(doc, [
    "基础通信：支持登录、注册、私聊、群聊、历史消息和未读提醒。",
    "群组协作：支持建群、邀请成员、公告维护、群主与管理员、禁言、退群和解散群。",
    "多媒体能力：支持文件上传、图片预览、文件下载、语音消息和局域网视频通话。",
    "产品体验：支持通知中心、个人资料、头像上传、主题色、聊天背景和移动端 APK。"
])
add_para(doc, "通过这些目标，我们希望项目不仅是一个页面展示，而是一个能够启动、登录、发送消息、保存数据，并且可以在手机上安装演示的系统。")

add_heading(doc, "三、总体架构设计")
add_para(doc, "系统架构采用前后端分离设计。前端使用 Vue 3 + Pinia + Vite，负责页面状态、聊天界面、个人中心、群管理面板、通知中心以及移动端响应式布局。后端按照 Spring Boot 的结构拆分 Controller、Service、Mapper/DAO，避免把所有逻辑写在一个文件中。")
add_para(doc, "实时通信部分从原生 WebSocket 逐步迁移为 Netty 4.1 长连接网关，用于处理在线状态、消息推送、心跳和 WebRTC 视频通话信令。数据层以 MySQL 保存用户、好友、群组、群成员、消息、通知和文件元数据；Redis 用于在线状态、未读数和会话令牌；MinIO 或对象存储用于保存真实文件内容，消息表只保存文件元数据和访问地址。")
add_callout(doc, "架构特点", "前端负责体验，后端负责业务规则，Netty 负责实时通道，MySQL/Redis/MinIO 分别承担持久化、状态缓存和文件存储。这样的拆分更符合设计报告中的后端架构，也便于后续扩展。", BLUE_SOFT)

add_heading(doc, "四、核心功能完成情况")
add_para(doc, "目前 FlowLink 的基础功能已经覆盖即时通讯平台的主要业务链路。账号体系方面，系统支持登录、注册、退出登录，并新增了“下次自动登录”的可选项，避免手机端一打开就进入上一次账号。个人资料支持昵称、邮箱、电话、地区、状态签名、头像等信息修改。")
add_para(doc, "聊天体验方面，系统支持私聊和群聊消息发送、历史消息分页加载、消息搜索高亮、引用回复、撤回和删除。消息栏和通讯录栏进行了职责区分，消息栏更关注会话，通讯录更关注联系人和群组。会话列表支持置顶和隐藏，右键菜单承担撤回、回复、删除等操作，让主界面更清爽。")
add_para(doc, "群聊方面，系统支持创建群聊、邀请成员、移除成员、设置管理员、转让群主、成员禁言、全员禁言、修改群公告和群介绍。近期我们重点修正了群昵称逻辑：群昵称只影响当前群聊消息栏和群成员列表，不影响通讯录、私聊和账号昵称。群主和管理员标识也已经显示在群成员列表和聊天消息中。")

add_heading(doc, "五、新增亮点功能")
add_para(doc, "第一是通知中心。好友申请、入群邀请、管理员变更、禁言通知等都集中展示，用户可以在通知中心快速查看和处理，避免信息散落在不同页面。")
add_para(doc, "第二是文件与媒体能力。系统支持文件上传、图片预览、文件下载、语音消息录制和播放。文件内容从前端 demo dataURL 迁移到对象存储思路，消息表和 file_record 表只保存元数据与访问地址，更接近真实工程。")
add_para(doc, "第三是局域网视频通话。我们采用 WebRTC 建立点对点音视频流，Netty WebSocket 只负责转发 call_invite、offer、answer、ice、hangup 等信令。这样服务端压力较小，适合局域网演示。移动端 APK 中还补充了相机、麦克风和音频设置权限，解决了 Android WebView 的 Permission denied 问题。")
add_para(doc, "第四是 AI 助手。系统中默认存在一个 AI 助手联系人，用户可以像私聊一样与其对话。后端通过 AiAssistantService 对接模型接口，复用消息链路生成 AI 回复，适合用于答辩问题准备、功能说明、Bug 分析和项目总结。")
add_para(doc, "第五是个性化体验。个人中心现在支持主题色选择、本地头像上传、聊天背景预设、自定义图片背景和退出登录。移动端输入栏按钮也进行了位置优化，让 APK 使用体验更接近真实聊天软件。")

add_heading(doc, "六、数据库与业务规则")
add_para(doc, "数据库设计围绕用户、好友、群组、群成员、消息、文件和通知展开。其中 group_member 表不仅保存用户是否在群内，还保存成员角色、群昵称、禁言状态等群内属性。这样可以把账号昵称和群昵称分开，保证同一个用户在不同群中可以有不同显示名。")
add_para(doc, "在业务规则上，我们重点避免了“只做前端效果”的问题。例如禁言不仅在界面上显示状态，后端发送消息时也会校验群成员是否被禁言、群聊是否全员禁言。文件上传后不会把内容直接塞进消息表，而是通过 file_record 保存文件名、大小、类型、访问地址等元数据。")

add_heading(doc, "七、测试验证与演示路径")
add_para(doc, "开发过程中，我们持续使用 npm build 验证前端构建，使用 Maven compile 验证后端代码，使用前后端联调验证消息发送、群管理和通知中心，使用 Android Studio 打包 APK 并在手机上测试权限、登录、退出和聊天布局。")
add_para(doc, "答辩演示建议按以下顺序进行：先展示登录和注册；然后进入私聊发送文本、图片或文件；接着进入群聊展示群昵称、群主/管理员标识、禁言和群公告；再展示通知中心的申请处理；随后演示语音消息和局域网视频通话；最后展示 AI 助手和个人中心的主题、头像、聊天背景设置。")

add_heading(doc, "八、项目难点")
add_bullets(doc, [
    "实时消息去重：发送方本地回显和 WebSocket 推送可能导致重复显示，需要通过 message id 或 clientId 合并。",
    "群昵称边界：群昵称只能作用于当前群聊，不能污染通讯录和私聊，这要求前端显示函数根据场景区分。",
    "禁言真实生效：禁言状态必须写入数据库并在后端发送接口校验，不能只停留在按钮状态。",
    "移动端权限：Android WebView 调用摄像头和麦克风时，需要同时处理 Manifest 权限、运行时授权和 WebView 权限请求。",
    "架构迁移：从简单 Demo 迁移到 Spring Boot Service、Mapper/DAO、Netty 网关和对象存储，需要重新划分职责边界。"
])

add_heading(doc, "九、项目管理与反思")
add_para(doc, "从开发管理角度看，这个项目经历了多次迭代：先完成基础 Demo，再逐步补齐数据库、后端服务、群聊管理、文件能力、通知中心、移动端 APK 和 AI 助手。每一轮迭代都围绕一个明确问题推进，例如“禁言没有实际效果”“消息刷新后丢失”“手机端无法退出登录”“群昵称影响范围不正确”等。")
add_para(doc, "项目优势在于功能覆盖较完整，前后端链路比较清晰，并且能够在浏览器和 Android 手机上进行真实演示。项目不足也比较明显：部分旧页面文本仍有待继续清理，自动化测试覆盖还不够充分，视频通话目前主要支持局域网一对一，生产环境还需要 TURN 中继、HTTPS 和更完善的部署方案。")
add_para(doc, "后续如果继续优化，我们会优先补充接口测试和单元测试，完善 Docker 一键部署脚本，加入消息加密和推送通知，并进一步扩展多人音视频会议。")

add_heading(doc, "十、结束语")
add_para(doc, "总的来说，FlowLink 从最初的文档设计和基础 Demo 出发，已经逐步演进为一个具备真实业务链路、后端分层架构、数据库持久化、实时通信、移动端 APK 和个性化体验的即时通讯系统。这个过程让我们更加理解了软件工程中需求分析、架构设计、数据库建模、前后端联调、测试验证和项目迭代之间的关系。")
add_para(doc, "以上就是我们小组的项目验收汇报，欢迎老师和同学提出问题。")

add_heading(doc, "附：答辩时可强调的项目特点")
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
headers = ["特点", "体现", "答辩表达"]
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    set_cell_fill(cell, MINT)
    set_cell_text(cell, h, bold=True, color=DARK_GREEN)

rows = [
    ("完整链路", "前端、后端、数据库、实时网关、对象存储、APK", "不是单纯页面原型，而是可启动、可登录、可保存、可演示的系统。"),
    ("真实 IM 体验", "右键菜单、群管理、通知中心、文件语音、未读提醒", "围绕真实聊天平台的高频操作进行设计。"),
    ("移动端能力", "Capacitor APK、权限处理、响应式布局", "项目可以从浏览器走向手机安装使用。"),
    ("技术扩展性", "Spring Boot Service、Netty、MySQL、Redis、MinIO、WebRTC", "后端结构清晰，后续可以继续扩展多人视频和生产部署。"),
    ("个性化亮点", "主题色、头像、本地上传背景、AI 助手", "让课程项目更接近一个有体验完整度的产品。"),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        set_cell_text(cells[i], text, size=9.5)
        if i == 0:
            set_cell_fill(cells[i], "F8FBF9")

for row in table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = Inches([1.35, 2.15, 2.9][idx])

doc.save(OUT)
print(OUT)
